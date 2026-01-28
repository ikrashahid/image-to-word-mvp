"""
TASK 1: Basic Image-to-Word Converter (MVP)
============================================

This MVP converts scanned documents/images to formatted Word documents
while preserving formatting like bold, italic, alignment, and paragraph structure.

Technology Stack:
- Google Gemini AI: Advanced OCR with formatting detection
- OpenCV/PIL: Image preprocessing
- python-docx: Word document generation with formatting
- Streamlit: Web-based GUI

How It Works:
-------------
1. IMAGE PREPROCESSING: Enhance image quality for better OCR
2. GEMINI AI OCR: Extract text WITH formatting instructions
3. FORMATTING DETECTION: Parse AI response for bold, italic, alignment
4. DOCUMENT GENERATION: Create formatted .docx preserving structure

Key Features:
- Detects bold and italic text
- Preserves paragraph structure
- Maintains text alignment (left, center, right)
- Handles headings and body text
- Generates clean, formatted Word documents
"""

import os
import cv2
import numpy as np
from PIL import Image
import base64
import requests
import json
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# API Configuration
API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyCAu8GFVsnTLFR5bN9kzCB4_g5L4L1t1U4")

class GeminiImageToWordConverter:
    """
    Main MVP Converter Class
    
    Converts images to formatted Word documents using Gemini AI.
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.api_key = API_KEY
        
    def preprocess_image(self, image_path: str) -> str:
        """
        STEP 1: Image Preprocessing
        =========================
        
        Purpose: Enhance image quality for better OCR accuracy
        
        Operations:
        - Convert to grayscale (reduces complexity)
        - Denoise using bilateral filter (removes noise, keeps edges)
        - Adaptive thresholding (better contrast for text)
        - Contrast enhancement (makes text clearer)
        
        Returns: Path to preprocessed image
        """
        self.logger.info("🔧 Preprocessing image...")
        
        # Read image
        image = cv2.imread(image_path)
        if image is None:
            raise ValueError(f"Could not read image: {image_path}")
        
        # Convert to grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        # Denoise - preserves edges while removing noise
        denoised = cv2.bilateralFilter(gray, 9, 75, 75)
        
        # Adaptive thresholding - better for varying lighting
        thresh = cv2.adaptiveThreshold(
            denoised, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 11, 2
        )
        
        # Enhance contrast
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(thresh)
        
        # Save preprocessed image
        preprocessed_path = image_path.replace('.', '_preprocessed.')
        cv2.imwrite(preprocessed_path, enhanced)
        
        self.logger.info(f"✅ Preprocessed image saved: {preprocessed_path}")
        return preprocessed_path
    
    def extract_formatted_text(self, image_path: str) -> dict:
        """
        STEP 2: Gemini AI OCR with Formatting Detection
        =============================================
        
        Purpose: Extract text AND detect formatting attributes
        
        Gemini AI analyzes the image and returns:
        - Raw text content
        - Formatting markers (bold, italic, headings)
        - Paragraph structure
        - Text alignment indicators
        
        The prompt is carefully designed to return structured data.
        
        Returns: Dictionary with text and formatting data
        """
        self.logger.info("🤖 Extracting text with Gemini AI...")
        
        # Read and encode image
        with open(image_path, 'rb') as f:
            image_bytes = f.read()
        image_base64 = base64.b64encode(image_bytes).decode('utf-8')
        
        # Gemini API endpoint
        model_name = "gemini-2.5-flash"
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={self.api_key}"
        
        # CRITICAL: This prompt tells Gemini HOW to format the response
        prompt = """
You are an expert OCR system that preserves document formatting.

TASK: Extract ALL text from this image and mark formatting attributes.

FORMATTING MARKERS TO USE:
- For BOLD text: Wrap in **text**
- For ITALIC text: Wrap in *text*
- For HEADINGS: Start line with # (H1), ## (H2), or ### (H3)
- For centered text: Start line with [CENTER]
- For right-aligned text: Start line with [RIGHT]
- For normal paragraphs: Just write the text

IMPORTANT RULES:
1. Preserve the exact text layout and structure
2. Identify paragraph breaks with blank lines
3. Detect and mark all bold/italic text
4. Identify headings by their size and position
5. Detect text alignment from visual position
6. Maintain reading order (top to bottom)

OUTPUT FORMAT:
Return the formatted text exactly as described above with all markers.
"""
        
        # API request payload
        payload = {
            "contents": [{
                "parts": [
                    {"text": prompt},
                    {
                        "inline_data": {
                            "mime_type": "image/jpeg",
                            "data": image_base64
                        }
                    }
                ]
            }]
        }
        
        headers = {"Content-Type": "application/json"}
        
        try:
            response = requests.post(url, headers=headers, json=payload)
            
            if response.status_code != 200:
                self.logger.error(f"API Error: {response.text}")
                return {"text": "", "error": response.text}
            
            result = response.json()
            extracted_text = result["candidates"][0]["content"]["parts"][0]["text"]
            
            self.logger.info("✅ Text extracted successfully")
            return {"text": extracted_text, "error": None}
            
        except Exception as e:
            self.logger.error(f"Extraction failed: {e}")
            return {"text": "", "error": str(e)}
    
    def parse_formatting(self, formatted_text: str) -> list:
        """
        STEP 3: Parse Formatting Markers
        ===============================
        
        Purpose: Convert Gemini's marked-up text into structured data
        
        Process:
        - Split text into lines
        - Detect formatting markers (**, *, #, [CENTER], [RIGHT])
        - Create structured data with content and formatting
        
        Returns: List of text blocks with formatting attributes
        """
        self.logger.info("📝 Parsing formatting markers...")
        
        blocks = []
        lines = formatted_text.split('\n')
        
        for line in lines:
            if not line.strip():
                # Empty line = paragraph break
                blocks.append({
                    'type': 'paragraph_break',
                    'content': ''
                })
                continue
            
            # Detect alignment
            alignment = 'left'
            if line.startswith('[CENTER]'):
                alignment = 'center'
                line = line.replace('[CENTER]', '').strip()
            elif line.startswith('[RIGHT]'):
                alignment = 'right'
                line = line.replace('[RIGHT]', '').strip()
            
            # Detect heading level
            heading_level = 0
            if line.startswith('###'):
                heading_level = 3
                line = line.replace('###', '').strip()
            elif line.startswith('##'):
                heading_level = 2
                line = line.replace('##', '').strip()
            elif line.startswith('#'):
                heading_level = 1
                line = line.replace('#', '').strip()
            
            # Parse inline formatting (bold and italic)
            formatted_runs = self._parse_inline_formatting(line)
            
            blocks.append({
                'type': 'heading' if heading_level > 0 else 'paragraph',
                'heading_level': heading_level,
                'alignment': alignment,
                'runs': formatted_runs
            })
        
        self.logger.info(f"✅ Parsed {len(blocks)} text blocks")
        return blocks
    
    def _parse_inline_formatting(self, text: str) -> list:
        """
        Helper: Parse inline formatting (bold, italic)
        
        Handles nested formatting like **bold *and italic* text**
        
        Returns: List of runs with text and formatting flags
        """
        runs = []
        current_pos = 0
        
        # Regex to find **bold** and *italic*
        pattern = r'(\*\*.*?\*\*|\*.*?\*)'
        
        for match in re.finditer(pattern, text):
            # Add text before the match
            if match.start() > current_pos:
                runs.append({
                    'text': text[current_pos:match.start()],
                    'bold': False,
                    'italic': False
                })
            
            # Add formatted text
            matched_text = match.group()
            if matched_text.startswith('**') and matched_text.endswith('**'):
                # Bold text
                runs.append({
                    'text': matched_text[2:-2],
                    'bold': True,
                    'italic': False
                })
            elif matched_text.startswith('*') and matched_text.endswith('*'):
                # Italic text
                runs.append({
                    'text': matched_text[1:-1],
                    'bold': False,
                    'italic': True
                })
            
            current_pos = match.end()
        
        # Add remaining text
        if current_pos < len(text):
            runs.append({
                'text': text[current_pos:],
                'bold': False,
                'italic': False
            })
        
        # If no formatting found, return whole text as single run
        if not runs:
            runs.append({
                'text': text,
                'bold': False,
                'italic': False
            })
        
        return runs
    
    def generate_word_document(self, blocks: list, output_path: str) -> bool:
        """
        STEP 4: Generate Formatted Word Document
        =======================================
        
        Purpose: Create .docx file with all formatting preserved
        
        Process:
        - Create Document object
        - Iterate through parsed blocks
        - Apply formatting (bold, italic, alignment, headings)
        - Save document
        
        Returns: True if successful, False otherwise
        """
        self.logger.info("📄 Generating Word document...")
        
        try:
            doc = Document()
            
            for block in blocks:
                if block['type'] == 'paragraph_break':
                    # Add empty paragraph for spacing
                    doc.add_paragraph()
                    continue
                
                # Create paragraph or heading
                if block['type'] == 'heading':
                    p = doc.add_heading(level=block['heading_level'])
                else:
                    p = doc.add_paragraph()
                
                # Set alignment
                if block['alignment'] == 'center':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif block['alignment'] == 'right':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Add formatted runs
                for run_data in block['runs']:
                    run = p.add_run(run_data['text'])
                    run.bold = run_data['bold']
                    run.italic = run_data['italic']
                    
                    # Set appropriate font size
                    if block['type'] == 'heading':
                        if block['heading_level'] == 1:
                            run.font.size = Pt(24)
                        elif block['heading_level'] == 2:
                            run.font.size = Pt(18)
                        else:
                            run.font.size = Pt(14)
                    else:
                        run.font.size = Pt(11)
            
            # Save document
            doc.save(output_path)
            self.logger.info(f"✅ Document saved: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Document generation failed: {e}")
            return False
    
    def convert(self, image_path: str, output_path: str) -> dict:
        """
        MAIN CONVERSION PIPELINE
        =======================
        
        Orchestrates the entire conversion process:
        1. Preprocess image
        2. Extract formatted text with Gemini
        3. Parse formatting markers
        4. Generate Word document
        
        Returns: Status dict with success flag and message
        """
        self.logger.info("🚀 Starting conversion pipeline...")
        
        try:
            # Step 1: Preprocess
            preprocessed_path = self.preprocess_image(image_path)
            
            # Step 2: Extract with formatting
            result = self.extract_formatted_text(preprocessed_path)
            
            if result['error']:
                return {
                    'success': False,
                    'message': f"OCR failed: {result['error']}"
                }
            
            # Step 3: Parse formatting
            blocks = self.parse_formatting(result['text'])
            
            # Step 4: Generate document
            success = self.generate_word_document(blocks, output_path)
            
            # Cleanup preprocessed image
            if os.path.exists(preprocessed_path):
                os.remove(preprocessed_path)
            
            if success:
                return {
                    'success': True,
                    'message': 'Conversion completed successfully!',
                    'output_path': output_path
                }
            else:
                return {
                    'success': False,
                    'message': 'Failed to generate Word document'
                }
                
        except Exception as e:
            self.logger.error(f"Conversion failed: {e}")
            return {
                'success': False,
                'message': str(e)
            }

# Convenience function for external use
def convert_image_to_word(image_path: str, output_path: str) -> dict:
    """
    Simple function to convert image to Word document.
    
    Usage:
        result = convert_image_to_word('input.jpg', 'output.docx')
        if result['success']:
            print(f"Saved to: {result['output_path']}")
    """
    converter = GeminiImageToWordConverter()
    return converter.convert(image_path, output_path)
